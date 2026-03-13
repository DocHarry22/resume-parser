"""Cover Letter Export Service - PDF, DOCX, and TXT generation."""

from io import BytesIO
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.cover_letter_models import CoverLetter, CoverLetterExportRequest


def generate_pdf(cover_letter: CoverLetter, request: CoverLetterExportRequest) -> bytes:
    """Generate PDF from cover letter."""
    buffer = BytesIO()
    
    # Set page size
    page_size = letter
    
    # Create document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Custom styles based on template
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
    )
    
    name_style = ParagraphStyle(
        'Name',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=6,
        textColor=colors.HexColor('#1a1a2e'),
    )
    
    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#666666'),
        spaceAfter=12,
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=12,
        firstLineIndent=0,
    )
    
    date_style = ParagraphStyle(
        'Date',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=24,
    )
    
    recipient_style = ParagraphStyle(
        'Recipient',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        spaceAfter=24,
    )
    
    salutation_style = ParagraphStyle(
        'Salutation',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=12,
    )
    
    closing_style = ParagraphStyle(
        'Closing',
        parent=styles['Normal'],
        fontSize=11,
        spaceBefore=24,
    )
    
    signature_style = ParagraphStyle(
        'Signature',
        parent=styles['Normal'],
        fontSize=11,
        spaceBefore=36,
    )
    
    # Build content
    story = []
    
    # Sender info (header)
    sender = cover_letter.sender
    story.append(Paragraph(sender.full_name or "Your Name", name_style))
    
    contact_parts = []
    if sender.email:
        contact_parts.append(sender.email)
    if sender.phone:
        contact_parts.append(sender.phone)
    if sender.linkedin:
        contact_parts.append(sender.linkedin)
    
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), contact_style))
    
    if sender.address or sender.city:
        address_parts = [sender.address, sender.city, sender.state, sender.zip_code]
        address = ", ".join(p for p in address_parts if p)
        if address:
            story.append(Paragraph(address, contact_style))
    
    story.append(Spacer(1, 24))
    
    # Date
    if cover_letter.settings.include_date:
        date_str = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(date_str, date_style))
    
    # Recipient info
    recipient = cover_letter.recipient
    recipient_lines = []
    if recipient.name:
        recipient_lines.append(recipient.name)
    if recipient.title:
        recipient_lines.append(recipient.title)
    if recipient.company:
        recipient_lines.append(recipient.company)
    if recipient.address:
        recipient_lines.append(recipient.address)
    city_state = ", ".join(p for p in [recipient.city, recipient.state, recipient.zip_code] if p)
    if city_state:
        recipient_lines.append(city_state)
    
    if recipient_lines:
        story.append(Paragraph("<br/>".join(recipient_lines), recipient_style))
    
    # Salutation
    salutation = cover_letter.content.salutation or "Dear Hiring Manager,"
    story.append(Paragraph(salutation, salutation_style))
    
    story.append(Spacer(1, 12))
    
    # Body content
    content = cover_letter.content
    
    if content.opening_paragraph:
        story.append(Paragraph(content.opening_paragraph, body_style))
    
    for paragraph in content.body_paragraphs:
        if paragraph:
            story.append(Paragraph(paragraph, body_style))
    
    if content.closing_paragraph:
        story.append(Paragraph(content.closing_paragraph, body_style))
    
    # Closing
    signature = content.signature or "Sincerely"
    story.append(Paragraph(f"{signature},", closing_style))
    story.append(Paragraph(cover_letter.sender.full_name or "Your Name", signature_style))
    
    # PS line
    if content.ps_line:
        story.append(Spacer(1, 24))
        story.append(Paragraph(f"P.S. {content.ps_line}", body_style))
    
    # Build PDF
    doc.build(story)
    
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx(cover_letter: CoverLetter, request: CoverLetterExportRequest) -> bytes:
    """Generate DOCX from cover letter."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Sender info (header)
    sender = cover_letter.sender
    
    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(sender.full_name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    # Contact info
    contact_parts = []
    if sender.email:
        contact_parts.append(sender.email)
    if sender.phone:
        contact_parts.append(sender.phone)
    if sender.linkedin:
        contact_parts.append(sender.linkedin)
    
    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.runs[0].font.size = Pt(10)
    
    if sender.address or sender.city:
        address_parts = [sender.address, sender.city, sender.state, sender.zip_code]
        address = ", ".join(p for p in address_parts if p)
        if address:
            addr_para = doc.add_paragraph(address)
            addr_para.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer
    
    # Date
    if cover_letter.settings.include_date:
        date_str = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(date_str)
        doc.add_paragraph()
    
    # Recipient info
    recipient = cover_letter.recipient
    if recipient.name:
        doc.add_paragraph(recipient.name)
    if recipient.title:
        doc.add_paragraph(recipient.title)
    if recipient.company:
        doc.add_paragraph(recipient.company)
    if recipient.address:
        doc.add_paragraph(recipient.address)
    city_state = ", ".join(p for p in [recipient.city, recipient.state, recipient.zip_code] if p)
    if city_state:
        doc.add_paragraph(city_state)
    
    doc.add_paragraph()  # Spacer
    
    # Salutation
    salutation = cover_letter.content.salutation or "Dear Hiring Manager,"
    doc.add_paragraph(salutation)
    
    doc.add_paragraph()  # Spacer
    
    # Body content
    content = cover_letter.content
    
    if content.opening_paragraph:
        doc.add_paragraph(content.opening_paragraph)
    
    for paragraph in content.body_paragraphs:
        if paragraph:
            doc.add_paragraph(paragraph)
    
    if content.closing_paragraph:
        doc.add_paragraph(content.closing_paragraph)
    
    doc.add_paragraph()  # Spacer
    
    # Closing
    signature = content.signature or "Sincerely"
    doc.add_paragraph(f"{signature},")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(cover_letter.sender.full_name or "Your Name")
    
    # PS line
    if content.ps_line:
        doc.add_paragraph()
        doc.add_paragraph(f"P.S. {content.ps_line}")
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_txt(cover_letter: CoverLetter) -> bytes:
    """Generate plain text from cover letter."""
    lines = []
    
    # Sender info
    sender = cover_letter.sender
    if sender.full_name:
        lines.append(sender.full_name)
    if sender.email or sender.phone:
        contact = " | ".join(p for p in [sender.email, sender.phone] if p)
        lines.append(contact)
    if sender.address or sender.city:
        address_parts = [sender.address, sender.city, sender.state, sender.zip_code]
        address = ", ".join(p for p in address_parts if p)
        if address:
            lines.append(address)
    
    lines.append("")
    
    # Date
    if cover_letter.settings.include_date:
        date_str = datetime.now().strftime("%B %d, %Y")
        lines.append(date_str)
        lines.append("")
    
    # Recipient info
    recipient = cover_letter.recipient
    if recipient.name:
        lines.append(recipient.name)
    if recipient.title:
        lines.append(recipient.title)
    if recipient.company:
        lines.append(recipient.company)
    if recipient.address:
        lines.append(recipient.address)
    city_state = ", ".join(p for p in [recipient.city, recipient.state, recipient.zip_code] if p)
    if city_state:
        lines.append(city_state)
    
    lines.append("")
    
    # Salutation
    salutation = cover_letter.content.salutation or "Dear Hiring Manager,"
    lines.append(salutation)
    
    lines.append("")
    
    # Body content
    content = cover_letter.content
    
    if content.opening_paragraph:
        lines.append(content.opening_paragraph)
        lines.append("")
    
    for paragraph in content.body_paragraphs:
        if paragraph:
            lines.append(paragraph)
            lines.append("")
    
    if content.closing_paragraph:
        lines.append(content.closing_paragraph)
        lines.append("")
    
    # Closing
    signature = content.signature or "Sincerely"
    lines.append(f"{signature},")
    lines.append("")
    lines.append(cover_letter.sender.full_name or "Your Name")
    
    # PS line
    if content.ps_line:
        lines.append("")
        lines.append(f"P.S. {content.ps_line}")
    
    return "\n".join(lines).encode('utf-8')
