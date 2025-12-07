"""DOCX document reader using python-docx.

DOCX files are much easier to parse than PDFs because they maintain
document structure (paragraphs, headings, tables, etc.) in XML format.
"""

from docx import Document
from pathlib import Path
from typing import BinaryIO, Union
from io import BytesIO
import re


class DOCXReader:
    """Extract text from DOCX files with structure preservation."""
    
    def __init__(self):
        """Initialize DOCX reader."""
        pass
    
    def read(self, file_source: Union[str, Path, BinaryIO, bytes]) -> dict:
        """
        Read a DOCX file and extract text content.
        
        Args:
            file_source: Path to DOCX file, file-like object, or bytes
            
        Returns:
            Dictionary with:
                - full_text: Complete text content
                - blocks: List of text blocks (paragraphs)
                - structured_blocks: Blocks with style information
        """
        # Handle bytes input by wrapping in BytesIO
        if isinstance(file_source, bytes):
            file_source = BytesIO(file_source)
        
        doc = Document(file_source)
        
        blocks = []
        structured_blocks = []
        full_text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Get paragraph style (for heading detection)
            style = paragraph.style.name if paragraph.style else 'Normal'
            
            blocks.append(text)
            structured_blocks.append({
                'text': text,
                'style': style,
                'is_heading': self._is_heading_style(style)
            })
            full_text_parts.append(text)
        
        # Process tables (if any)
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                blocks.append(table_text)
                structured_blocks.append({
                    'text': table_text,
                    'style': 'Table',
                    'is_heading': False
                })
                full_text_parts.append(table_text)
        
        full_text = '\n\n'.join(full_text_parts)
        
        return {
            'full_text': full_text,
            'blocks': blocks,
            'structured_blocks': structured_blocks
        }
    
    def _is_heading_style(self, style: str) -> bool:
        """
        Check if a paragraph style is a heading.
        
        Args:
            style: Paragraph style name
            
        Returns:
            True if style is a heading
        """
        style_lower = style.lower()
        
        # Common heading styles
        heading_patterns = [
            r'^heading\s*\d*$',
            r'^title$',
            r'^subtitle$',
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, style_lower):
                return True
        
        return False
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: Table object from python-docx
            
        Returns:
            Formatted table text
        """
        rows_text = []
        
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells_text.append(cell_text)
            
            if cells_text:
                rows_text.append(' | '.join(cells_text))
        
        return '\n'.join(rows_text) if rows_text else ''
